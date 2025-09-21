"""
Text extraction utilities for different file formats.

This module provides utilities for extracting text from various resume formats
including PDF, DOCX, and TXT files with proper error handling and formatting.
"""

import os
import re
import logging
from typing import Optional, Dict, Any, List
from pathlib import Path

# PDF processing
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

try:
    import fitz  # PyMuPDF
    FITZ_AVAILABLE = True
except ImportError:
    FITZ_AVAILABLE = False

# DOCX processing
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import docx2txt
    DOCX2TXT_AVAILABLE = True
except ImportError:
    DOCX2TXT_AVAILABLE = False

from config import config


class TextExtractor:
    """Handles text extraction from various file formats."""
    
    def __init__(self):
        """Initialize the text extractor."""
        self.logger = logging.getLogger(__name__)
        self.config = config.parsing
        
        # Check available libraries
        self._check_dependencies()
    
    def _check_dependencies(self):
        """Check which parsing libraries are available."""
        missing_libs = []
        
        if not PDFPLUMBER_AVAILABLE and not PYPDF2_AVAILABLE:
            missing_libs.append("pdfplumber or PyPDF2 for PDF parsing")
        
        if not DOCX_AVAILABLE and not DOCX2TXT_AVAILABLE:
            missing_libs.append("python-docx or docx2txt for DOCX parsing")
        
        if missing_libs:
            self.logger.warning(f"Missing optional dependencies: {', '.join(missing_libs)}")
    
    def extract_text(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from a file based on its format.
        
        Args:
            file_path: Path to the file to extract text from
            
        Returns:
            Dictionary containing extracted text and metadata
        """
        result = {
            "text": "",
            "metadata": {},
            "errors": [],
            "success": False
        }
        
        try:
            # Validate file
            if not os.path.exists(file_path):
                result["errors"].append(f"File not found: {file_path}")
                return result
            
            file_path = Path(file_path)
            file_size_mb = file_path.stat().st_size / (1024 * 1024)
            
            # Check file size
            if file_size_mb > self.config.max_file_size_mb:
                result["errors"].append(
                    f"File size ({file_size_mb:.1f}MB) exceeds limit ({self.config.max_file_size_mb}MB)"
                )
                return result
            
            # Determine file type and extract text
            file_extension = file_path.suffix.lower()
            
            if file_extension == ".pdf":
                result = self._extract_pdf_text(file_path)
            elif file_extension in [".docx", ".doc"]:
                result = self._extract_docx_text(file_path)
            elif file_extension == ".txt":
                result = self._extract_txt_text(file_path)
            else:
                result["errors"].append(f"Unsupported file format: {file_extension}")
                return result
            
            # Add common metadata
            result["metadata"].update({
                "file_name": file_path.name,
                "file_size_mb": round(file_size_mb, 2),
                "file_extension": file_extension
            })
            
            # Post-process text if extraction was successful
            if result["success"] and result["text"]:
                result["text"] = self._clean_text(result["text"])
                result["metadata"]["text_length"] = len(result["text"])
                result["metadata"]["word_count"] = len(result["text"].split())
            
        except Exception as e:
            self.logger.error(f"Error extracting text from {file_path}: {str(e)}")
            result["errors"].append(f"Extraction error: {str(e)}")
        
        return result
    
    def _extract_pdf_text(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from PDF file."""
        result = {
            "text": "",
            "metadata": {"page_count": 0, "extraction_method": ""},
            "errors": [],
            "success": False
        }
        
        # Try PyMuPDF first (fastest and most reliable)
        if FITZ_AVAILABLE:
            try:
                doc = fitz.open(file_path)
                text_parts = []
                result["metadata"]["page_count"] = doc.page_count
                result["metadata"]["extraction_method"] = "PyMuPDF"
                
                for page_num in range(doc.page_count):
                    try:
                        page = doc[page_num]
                        page_text = page.get_text()
                        if page_text:
                            text_parts.append(page_text)
                    except Exception as e:
                        self.logger.warning(f"Error extracting page {page_num + 1}: {str(e)}")
                        result["errors"].append(f"Page {page_num + 1} extraction error")
                
                doc.close()
                result["text"] = "\n".join(text_parts)
                result["success"] = True
                return result
                
            except Exception as e:
                self.logger.warning(f"PyMuPDF failed: {str(e)}")
                result["errors"].append(f"PyMuPDF error: {str(e)}")
        
        # Try pdfplumber second (better for complex layouts)
        if PDFPLUMBER_AVAILABLE:
            try:
                with pdfplumber.open(file_path) as pdf:
                    text_parts = []
                    result["metadata"]["page_count"] = len(pdf.pages)
                    result["metadata"]["extraction_method"] = "pdfplumber"
                    
                    for page_num, page in enumerate(pdf.pages):
                        try:
                            page_text = page.extract_text()
                            if page_text:
                                text_parts.append(page_text)
                            
                            # Extract tables if present
                            tables = page.extract_tables()
                            for table in tables:
                                if table:
                                    table_text = self._table_to_text(table)
                                    text_parts.append(table_text)
                                    
                        except Exception as e:
                            self.logger.warning(f"Error extracting page {page_num + 1}: {str(e)}")
                            result["errors"].append(f"Page {page_num + 1} extraction error")
                    
                    result["text"] = "\n".join(text_parts)
                    result["success"] = True
                    return result
                    
            except Exception as e:
                self.logger.warning(f"pdfplumber failed: {str(e)}")
                result["errors"].append(f"pdfplumber error: {str(e)}")
        
        # Fallback to PyPDF2
        if PYPDF2_AVAILABLE:
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    text_parts = []
                    result["metadata"]["page_count"] = len(pdf_reader.pages)
                    result["metadata"]["extraction_method"] = "PyPDF2"
                    
                    for page_num, page in enumerate(pdf_reader.pages):
                        try:
                            page_text = page.extract_text()
                            if page_text:
                                text_parts.append(page_text)
                        except Exception as e:
                            self.logger.warning(f"Error extracting page {page_num + 1}: {str(e)}")
                            result["errors"].append(f"Page {page_num + 1} extraction error")
                    
                    result["text"] = "\n".join(text_parts)
                    result["success"] = True
                    return result
                    
            except Exception as e:
                self.logger.error(f"PyPDF2 failed: {str(e)}")
                result["errors"].append(f"PyPDF2 error: {str(e)}")
        
        if not result["success"]:
            result["errors"].append("No PDF parsing library available or all methods failed")
        
        return result
    
    def _extract_docx_text(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from DOCX file."""
        result = {
            "text": "",
            "metadata": {"extraction_method": ""},
            "errors": [],
            "success": False
        }
        
        # Try python-docx first (more detailed extraction)
        if DOCX_AVAILABLE:
            try:
                doc = Document(file_path)
                text_parts = []
                result["metadata"]["extraction_method"] = "python-docx"
                
                # Extract paragraph text
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text)
                
                # Extract table text
                for table in doc.tables:
                    table_text = self._docx_table_to_text(table)
                    if table_text:
                        text_parts.append(table_text)
                
                result["text"] = "\n".join(text_parts)
                result["success"] = True
                return result
                
            except Exception as e:
                self.logger.warning(f"python-docx failed: {str(e)}")
                result["errors"].append(f"python-docx error: {str(e)}")
        
        # Fallback to docx2txt
        if DOCX2TXT_AVAILABLE:
            try:
                text = docx2txt.process(str(file_path))
                result["text"] = text
                result["metadata"]["extraction_method"] = "docx2txt"
                result["success"] = True
                return result
                
            except Exception as e:
                self.logger.error(f"docx2txt failed: {str(e)}")
                result["errors"].append(f"docx2txt error: {str(e)}")
        
        if not result["success"]:
            result["errors"].append("No DOCX parsing library available or all methods failed")
        
        return result
    
    def _extract_txt_text(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from TXT file."""
        result = {
            "text": "",
            "metadata": {"encoding": ""},
            "errors": [],
            "success": False
        }
        
        # Try different encodings
        encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    text = file.read()
                    result["text"] = text
                    result["metadata"]["encoding"] = encoding
                    result["success"] = True
                    return result
                    
            except UnicodeDecodeError:
                continue
            except Exception as e:
                result["errors"].append(f"Error with encoding {encoding}: {str(e)}")
        
        result["errors"].append("Could not decode file with any supported encoding")
        return result
    
    def _table_to_text(self, table: List[List[str]]) -> str:
        """Convert pdfplumber table to text."""
        if not table:
            return ""
        
        text_rows = []
        for row in table:
            if row:
                # Filter out None values and convert to strings
                clean_row = [str(cell) if cell is not None else "" for cell in row]
                if any(cell.strip() for cell in clean_row):  # Only add non-empty rows
                    text_rows.append(" | ".join(clean_row))
        
        return "\n".join(text_rows)
    
    def _docx_table_to_text(self, table) -> str:
        """Convert python-docx table to text."""
        text_rows = []
        
        for row in table.rows:
            row_cells = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_cells.append(cell_text)
            
            if row_cells:
                text_rows.append(" | ".join(row_cells))
        
        return "\n".join(text_rows)
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text."""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove common resume artifacts
        text = re.sub(r'Page \d+ of \d+', '', text, flags=re.IGNORECASE)
        text = re.sub(r'Page \d+', '', text, flags=re.IGNORECASE)
        
        # Remove extra line breaks but preserve paragraph structure
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        
        # Remove leading/trailing whitespace
        text = text.strip()
        
        # Limit text length if specified
        if hasattr(self.config, 'max_text_length') and len(text) > self.config.max_text_length:
            text = text[:self.config.max_text_length]
            self.logger.warning(f"Text truncated to {self.config.max_text_length} characters")
        
        return text
    
    def extract_contact_info(self, text: str) -> Dict[str, Optional[str]]:
        """
        Extract contact information from text using regex patterns.
        
        Args:
            text: Text to extract contact info from
            
        Returns:
            Dictionary with extracted contact information
        """
        contact_info = {
            "email": None,
            "phone": None,
            "linkedin": None,
            "github": None
        }
        
        try:
            # Email extraction
            email_match = re.search(self.config.email_pattern, text, re.IGNORECASE)
            if email_match:
                contact_info["email"] = email_match.group()
            
            # Phone extraction
            phone_match = re.search(self.config.phone_pattern, text)
            if phone_match:
                contact_info["phone"] = phone_match.group()
            
            # LinkedIn extraction
            linkedin_patterns = [
                r'linkedin\.com/in/([a-zA-Z0-9-]+)',
                r'linkedin\.com/pub/([a-zA-Z0-9-]+)',
                r'linkedin\.com/profile/view\?id=([a-zA-Z0-9-]+)'
            ]
            
            for pattern in linkedin_patterns:
                match = re.search(pattern, text, re.IGNORECASE)
                if match:
                    contact_info["linkedin"] = match.group()
                    break
            
            # GitHub extraction
            github_patterns = [
                r'github\.com/([a-zA-Z0-9-]+)',
                r'github\.io/([a-zA-Z0-9-]+)'
            ]
            
            for pattern in github_patterns:
                match = re.search(pattern, text, re.IGNORECASE)
                if match:
                    contact_info["github"] = match.group()
                    break
                    
        except Exception as e:
            self.logger.error(f"Error extracting contact info: {str(e)}")
        
        return contact_info
    
    def validate_extraction(self, result: Dict[str, Any]) -> bool:
        """
        Validate that text extraction was successful and meaningful.
        
        Args:
            result: Extraction result dictionary
            
        Returns:
            True if extraction is valid
        """
        if not result["success"]:
            return False
        
        text = result["text"]
        if not text or len(text.strip()) < 50:  # Minimum meaningful content
            return False
        
        # Check for reasonable word count
        word_count = len(text.split())
        if word_count < 10:
            return False
        
        # Check for presence of likely resume content
        resume_indicators = [
            'experience', 'education', 'skills', 'work', 'employment',
            'university', 'college', 'degree', 'project', 'responsibility'
        ]
        
        text_lower = text.lower()
        indicator_count = sum(1 for indicator in resume_indicators if indicator in text_lower)
        
        # Should have at least 2 resume indicators
        return indicator_count >= 2

